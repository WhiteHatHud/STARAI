from pathlib import Path
import tempfile
import os
import re
import json
import pandas as pd
from typing import Dict, Any, Optional, List
from jinja2 import Environment, FileSystemLoader
from bson import ObjectId
from docx import Document
from docx.shared import Inches, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.shared import OxmlElement, qn
from openpyxl import Workbook
from openpyxl.styles import Font, PatternFill, Alignment
from openpyxl.utils.dataframe import dataframe_to_rows

class DocumentFormatter:
    """Helper class for document formatting operations"""
    
    @staticmethod
    def get_alignment_from_formatting(formatting: List[str]) -> str:
        """Get alignment value from formatting array"""
        if 'center' in formatting:
            return 'center'
        elif 'right' in formatting:
            return 'right'
        else:
            return 'left'
    
    @staticmethod
    def get_docx_alignment(formatting: List[str]):
        """Get DOCX alignment constant from formatting array"""
        if 'center' in formatting:
            return WD_ALIGN_PARAGRAPH.CENTER
        elif 'right' in formatting:
            return WD_ALIGN_PARAGRAPH.RIGHT
        else:
            return WD_ALIGN_PARAGRAPH.LEFT
    
    @staticmethod
    def apply_text_formatting_to_run(run, formatting: List[str]):
        """Apply text formatting (bold, italic, etc.) to a DOCX run"""
        if 'bold' in formatting:
            run.bold = True
        if 'italic' in formatting:
            run.italic = True
        if 'underline' in formatting:
            run.underline = True
        if 'strikethrough' in formatting:
            run.font.strike = True
    
class DocxHelper:
    """Helper class for DOCX operations"""
    
    @staticmethod
    def setup_document_defaults(doc: Document) -> Document:
        """Set up default document styling"""
        style = doc.styles['Normal']
        font = style.font
        font.name = 'Times New Roman'
        font.size = Pt(12)
        
        paragraph_format = style.paragraph_format
        paragraph_format.line_spacing = 1.3
        paragraph_format.space_after = Pt(6)
        
        return doc
    
    @staticmethod
    def create_paragraph_with_formatting(doc: Document, content: str, formatting: List[str], 
                                       space_before: int = 0, space_after: int = 6) -> None:
        """Create a paragraph with specified formatting"""
        paragraph = doc.add_paragraph()
        paragraph.paragraph_format.space_before = Pt(space_before)
        paragraph.paragraph_format.space_after = Pt(space_after)
        paragraph.alignment = DocumentFormatter.get_docx_alignment(formatting)
        
        run = paragraph.add_run(content)
        run.font.name = 'Times New Roman'
        run.font.size = Pt(12)
        DocumentFormatter.apply_text_formatting_to_run(run, formatting)
    
    @staticmethod
    def create_section_title(doc: Document, title: str) -> None:
        """Create a section title with proper formatting"""
        section_title_paragraph = doc.add_paragraph()
        section_title_paragraph.paragraph_format.space_before = Pt(10)
        section_title_paragraph.paragraph_format.space_after = Pt(6)
        section_title_paragraph.paragraph_format.border_bottom = True
        
        section_title_run = section_title_paragraph.add_run(title)
        section_title_run.font.name = 'Times New Roman'
        section_title_run.font.size = Pt(12)
        section_title_run.bold = True
    
    @staticmethod
    def create_list_with_formatting(doc: Document, items: List[str], formatting: List[str]) -> None:
        """Create a list with specified formatting"""
        for item in items:
            list_paragraph = doc.add_paragraph()
            list_paragraph.style = 'List Bullet'
            list_paragraph.alignment = DocumentFormatter.get_docx_alignment(formatting)
            list_paragraph.paragraph_format.space_after = Pt(3)
            list_paragraph.paragraph_format.left_indent = Inches(0.25)
            
            list_run = list_paragraph.add_run(item)
            list_run.font.name = 'Times New Roman'
            list_run.font.size = Pt(12)
            DocumentFormatter.apply_text_formatting_to_run(list_run, formatting)
    
    @staticmethod
    def create_table_with_formatting(doc: Document, table_data: List[List[str]], 
                                   formatting: List[str]) -> None:
        """Create a table with specified formatting"""
        if not table_data or len(table_data) == 0:
            return
        
        # Create a paragraph for table positioning if center/right aligned
        if 'center' in formatting or 'right' in formatting:
            table_positioning_paragraph = doc.add_paragraph()
            table_positioning_paragraph.alignment = DocumentFormatter.get_docx_alignment(formatting)
            table_positioning_paragraph.paragraph_format.space_after = Pt(0)
        
        # Create a separate table for each row to allow different column counts
        for row_idx, row_data in enumerate(table_data):
            if not row_data:
                continue
                
            num_cols = len(row_data)
            table = doc.add_table(rows=1, cols=num_cols)
            table.style = 'Table Grid'
            table.autofit = False
            table.allow_autofit = False
            
            # Set table alignment based on formatting
            table.alignment = DocumentFormatter.get_docx_alignment(formatting)
            
            # Calculate column widths
            total_width = Inches(6.5)
            col_width = total_width / num_cols
            
            # Set table borders
            TableHelper.set_table_borders(table)
            
            # Fill the row with data
            for col_idx, cell_data in enumerate(row_data):
                cell = table.cell(0, col_idx)
                cell.width = col_width
                cell_paragraph = cell.paragraphs[0]
                cell_paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
                cell_paragraph.paragraph_format.space_after = Pt(0)
                cell_paragraph.clear()
                
                cell_run = cell_paragraph.add_run(str(cell_data))
                cell_run.font.name = 'Times New Roman'
                cell_run.font.size = Pt(9)
                DocumentFormatter.apply_text_formatting_to_run(cell_run, formatting)
        
        # Add spacing after tables
        table_spacer = doc.add_paragraph()
        table_spacer.paragraph_format.space_after = Pt(8)

class TableHelper:
    """Helper class for table operations"""
    
    @staticmethod
    def set_table_borders(table):
        """Set table borders for DOCX table"""
        from docx.oxml.ns import nsdecls, qn
        from docx.oxml import parse_xml
        
        # Get table element
        tbl = table._tbl
        
        # Create table borders element
        tbl_borders = parse_xml(f'''
            <w:tblBorders {nsdecls('w')}>
                <w:top w:val="single" w:sz="4" w:space="0" w:color="000000"/>
                <w:left w:val="single" w:sz="4" w:space="0" w:color="000000"/>
                <w:bottom w:val="single" w:sz="4" w:space="0" w:color="000000"/>
                <w:right w:val="single" w:sz="4" w:space="0" w:color="000000"/>
                <w:insideH w:val="single" w:sz="4" w:space="0" w:color="000000"/>
                <w:insideV w:val="single" w:sz="4" w:space="0" w:color="000000"/>
            </w:tblBorders>
        ''')
        
        # Get or create table properties
        tbl_pr = tbl.find(qn('w:tblPr'))
        if tbl_pr is None:
            tbl_pr = OxmlElement('w:tblPr')
            tbl.insert(0, tbl_pr)
        
        # Remove existing borders if any
        existing_borders = tbl_pr.find(qn('w:tblBorders'))
        if existing_borders is not None:
            tbl_pr.remove(existing_borders)
        
        # Add new borders
        tbl_pr.append(tbl_borders)

class TemplateFilters:
    """Static filters for Jinja2 templates"""
    
    @staticmethod
    def chr_offset(value):
        """Convert number to letter (0=a, 1=b, etc.)"""
        return chr(ord('a') + value)

    @staticmethod
    def nl2br(value):
        """Convert newlines to HTML breaks with proper spacing"""
        if not value:
            return ""
        # Replace double newlines (paragraph breaks) with </p><p>
        value = re.sub(r'\n\s*\n', '</p><p>', str(value))
        # Replace single newlines with <br>
        value = re.sub(r'([^\n])\n([^\n])', r'\1<br>\2', value)
        # Wrap in paragraph tags if not already wrapped
        if not value.startswith('<p>'):
            value = f'<p>{value}</p>'
        return value
    
    @staticmethod
    def clean_br_tags(value):
        """Clean up excessive br tags and spacing"""
        if not value:
            return ""
        # Remove multiple consecutive br tags
        value = re.sub(r'(<br\s*/?>\s*){3,}', '<br><br>', str(value))
        # Remove br tags at the start and end of paragraphs
        value = re.sub(r'<p>\s*<br\s*/?>', '<p>', value)
        value = re.sub(r'<br\s*/?>\s*</p>', '</p>', value)
        # Clean up whitespace around br tags
        value = re.sub(r'\s*<br\s*/?>\s*', '<br>', value)
        return value

class CustomReportReportService:
    """Main service class for generating custom case study reports"""
    
    # Setup Jinja2 environment for templates
    template_dir = Path(__file__).parent.parent / "templates"
    env = Environment(loader=FileSystemLoader(str(template_dir)))
    env.filters['chr_offset'] = TemplateFilters.chr_offset
    env.filters['nl2br'] = TemplateFilters.nl2br
    env.filters['clean_br_tags'] = TemplateFilters.clean_br_tags
    
    @staticmethod
    def _prepare_report_data(report_data: Dict[str, Any]) -> Dict[str, Any]:
        """Prepare case study data by converting Pydantic models if necessary"""
        if hasattr(report_data, "dict"):
            return report_data.dict()
        elif hasattr(report_data, "model_dump"):
            return report_data.model_dump()
        return report_data
    
    @staticmethod
    def _generate_output_path(prefix: str, extension: str, output_path: Optional[str] = None) -> str:
        """Generate output path if not provided"""
        if not output_path:
            temp_dir = tempfile.gettempdir()
            output_path = os.path.join(temp_dir, f"{prefix}_{ObjectId()}.{extension}")
        return output_path
    
    @staticmethod
    def generate_pdf_custom_report(report_data: Dict[str, Any], output_path: Optional[str] = None) -> str:
        """Generate PDF using LibreOffice conversion from DOCX"""
        import tempfile
        import subprocess
        import shutil
        
        temp_docx_path = None
        try:
            with tempfile.NamedTemporaryFile(suffix='.docx', delete=False) as temp_docx:
                temp_docx_path = temp_docx.name
            
            CustomReportReportService.generate_docx_custom_report(
                report_data, temp_docx_path
            )
            
            if not output_path:
                output_path = CustomReportReportService._generate_output_path(
                    "custom_report", "pdf", None
                )
            
            output_dir = os.path.dirname(output_path)
            
            print("Converting DOCX to PDF using LibreOffice...")
            
            result = subprocess.run([
                'soffice',
                '--headless',
                '--convert-to', 'pdf',
                '--outdir', output_dir,
                temp_docx_path
            ], 
            capture_output=True, 
            text=True,
            timeout=60 
            )
            
            if result.returncode != 0:
                raise RuntimeError(f"LibreOffice conversion failed: {result.stderr}")
            
            temp_docx_name = os.path.basename(temp_docx_path)
            temp_pdf_name = temp_docx_name.replace('.docx', '.pdf')
            temp_pdf_path = os.path.join(output_dir, temp_pdf_name)
            
            if temp_pdf_path != output_path:
                shutil.move(temp_pdf_path, output_path)
            
            if not os.path.exists(output_path):
                raise RuntimeError("PDF file was not created successfully")
            
            print(f"✅ PDF generated successfully using LibreOffice: {output_path}")
            return output_path
            
        except subprocess.TimeoutExpired:
            raise RuntimeError("LibreOffice conversion timed out after 60 seconds")
        except FileNotFoundError:
            raise RuntimeError("LibreOffice not found. Please install LibreOffice.")
        except Exception as e:
            raise RuntimeError(f"LibreOffice conversion failed: {str(e)}")
        finally:
            if temp_docx_path and os.path.exists(temp_docx_path):
                os.unlink(temp_docx_path)

    @staticmethod
    def generate_docx_custom_report(report_data: Dict[str, Any], output_path: Optional[str] = None):
        """Generate a DOCX from custom format case study data with proper text, list, and table rendering"""
        
        # Prepare data
        report_data = CustomReportReportService._prepare_report_data(report_data)
        sections = report_data.get("sections", [])
        header = report_data.get("header", {})
        footer = report_data.get("footer", {})

        # Create and setup document
        doc = Document()
        DocxHelper.setup_document_defaults(doc)
        
        # Add header if configured
        if header and header.get("type") and header["type"] != "":
            CustomReportReportService._add_header(doc, header)
        
        # Add footer if configured
        if footer and footer.get("type") and footer["type"] != "":
            CustomReportReportService._add_footer(doc, footer)
        
        # Process sections
        for section in sections:
            section_content = section.get("content", "")
            content_type = section.get("content_type", "")
            formatting = section.get("formatting", [])
            
            # Parse JSON content
            parsed_content = json.loads(section_content) if isinstance(section_content, str) else section_content
            
            # Handle different content types
            if content_type == 'structural':
                if parsed_content.get('textdata'):
                    DocxHelper.create_paragraph_with_formatting(
                        doc, parsed_content['textdata'], formatting, 
                        space_before=12, space_after=8
                    )
            
            elif content_type == 'formatting':
                if parsed_content.get('textdata'):
                    DocxHelper.create_paragraph_with_formatting(
                        doc, parsed_content['textdata'], formatting, 
                        space_before=8, space_after=6
                    )
            
            elif content_type == 'content':
                # Process content based on type
                if parsed_content.get('textdata'):
                    DocxHelper.create_paragraph_with_formatting(
                        doc, parsed_content['textdata'], formatting
                    )
                elif parsed_content.get('listdata'):
                    DocxHelper.create_list_with_formatting(
                        doc, parsed_content['listdata'], formatting
                    )
                elif parsed_content.get('tabledata'):
                    DocxHelper.create_table_with_formatting(
                        doc, parsed_content['tabledata'], formatting
                    )
        
        # Generate output path and save
        output_path = CustomReportReportService._generate_output_path(
            "custom_report", "docx", output_path
        )
        doc.save(output_path)
        return output_path
    
    @staticmethod
    def _add_header(doc: Document, header_config: Dict[str, Any]):
        """Add header to the document"""
        header_type = header_config.get("type", "")
        header_formatting = header_config.get("text_formatting", [])
        
        # Get the header section
        section = doc.sections[0]
        header = section.header
        
        # Create header paragraph
        header_paragraph = header.paragraphs[0]
        header_paragraph.alignment = DocumentFormatter.get_docx_alignment(header_formatting)
        
        if header_type == "text":
            # Add text content
            content = header_config.get("content", "")
            if content:
                run = header_paragraph.add_run(content)
                run.font.name = 'Times New Roman'
                run.font.size = Pt(10)
                DocumentFormatter.apply_text_formatting_to_run(run, header_formatting)
        
        elif header_type == "page_number":
            # Add page number
            run = header_paragraph.add_run("Page ")
            run.font.name = 'Times New Roman'
            run.font.size = Pt(10)
            DocumentFormatter.apply_text_formatting_to_run(run, header_formatting)
            
            fldChar1 = OxmlElement('w:fldChar')
            fldChar1.set(qn('w:fldCharType'), 'begin')
            
            instrText = OxmlElement('w:instrText')
            instrText.text = "PAGE"
            
            fldChar2 = OxmlElement('w:fldChar')
            fldChar2.set(qn('w:fldCharType'), 'end')
            
            run._r.append(fldChar1)
            run._r.append(instrText)
            run._r.append(fldChar2)
    
    @staticmethod
    def _add_footer(doc: Document, footer_config: Dict[str, Any]):
        """Add footer to the document"""
        footer_type = footer_config.get("type", "")
        footer_formatting = footer_config.get("text_formatting", [])
        
        # Get the footer section
        section = doc.sections[0]
        footer = section.footer
        
        # Create footer paragraph
        footer_paragraph = footer.paragraphs[0]
        footer_paragraph.alignment = DocumentFormatter.get_docx_alignment(footer_formatting)
        
        if footer_type == "text":
            # Add text content
            content = footer_config.get("content", "")
            if content:
                run = footer_paragraph.add_run(content)
                run.font.name = 'Times New Roman'
                run.font.size = Pt(10)
                DocumentFormatter.apply_text_formatting_to_run(run, footer_formatting)
        
        elif footer_type == "page_number":
            # Add page number
            run = footer_paragraph.add_run("Page ")
            run.font.name = 'Times New Roman'
            run.font.size = Pt(10)
            DocumentFormatter.apply_text_formatting_to_run(run, footer_formatting)
            
            fldChar1 = OxmlElement('w:fldChar')
            fldChar1.set(qn('w:fldCharType'), 'begin')
            
            instrText = OxmlElement('w:instrText')
            instrText.text = "PAGE"
            
            fldChar2 = OxmlElement('w:fldChar')
            fldChar2.set(qn('w:fldCharType'), 'end')
            
            run._r.append(fldChar1)
            run._r.append(instrText)
            run._r.append(fldChar2)
    
    @staticmethod
    def _set_table_borders(table):
        """Backward compatibility wrapper"""
        return TableHelper.set_table_borders(table)
    
    @staticmethod
    def format_content_for_excel(content) -> str:
        """Format content for Excel cells, handling textdata, listdata, and tabledata structures"""
        if not content:
            return ""
        
        text = str(content)
        
        try:
            parsed = json.loads(text)
            
            if isinstance(parsed, dict):
                if 'textdata' in parsed:
                    text_content = parsed['textdata']
                    if isinstance(text_content, str):
                        text_content = re.sub(r'<[^>]+>', '', text_content)
                        text_content = re.sub(r'\n\s*\n\s*\n', '\n\n', text_content)
                        text_content = re.sub(r'[ \t]+', ' ', text_content)
                        return text_content.strip()
                    return str(text_content)
                
                elif 'listdata' in parsed:
                    list_content = parsed['listdata']
                    if isinstance(list_content, list):
                        formatted_items = []
                        for item in list_content:
                            item_text = str(item).strip()
                            if item_text:
                                formatted_items.append(f"• {item_text}")
                        return "\n".join(formatted_items)
                    return str(list_content)
                
                elif 'tabledata' in parsed:
                    table_content = parsed['tabledata']
                    if isinstance(table_content, list) and table_content:
                        formatted_rows = []
                        
                        for row_idx, row in enumerate(table_content):
                            if isinstance(row, list):
                                clean_cells = []
                                for cell in row:
                                    cell_text = str(cell).strip() if cell is not None else ""
                                    cell_text = re.sub(r'<[^>]+>', '', cell_text)
                                    clean_cells.append(cell_text)
                                
                                if any(cell.strip() for cell in clean_cells):
                                    row_text = " | ".join(clean_cells)
                                    
                                    if row_idx == 0:
                                        formatted_rows.append(f"[HEADER] {row_text}")
                                        formatted_rows.append("-" * 119)
                                    else:
                                        formatted_rows.append(f"[ROW {row_idx}] {row_text}")
                                        formatted_rows.append("-" * 119)
                        
                        if formatted_rows:
                            return "\n".join(formatted_rows)
                    return str(table_content)
            
            # Simple fallback for any other JSON structure
            return json.dumps(parsed, indent=2) if isinstance(parsed, (dict, list)) else str(parsed)
                
        except (json.JSONDecodeError, Exception):
            pass
        
        # Clean up as regular text
        text = re.sub(r'<[^>]+>', '', text)
        text = re.sub(r'\n\s*\n\s*\n', '\n\n', text)
        text = re.sub(r'[ \t]+', ' ', text)
        return text.strip()

    def extract_section_content(sections) -> dict:
        """Extract section headings and content from case study"""
        content = {}
        if not sections:
            return content
            
        for section in sections:
            title = getattr(section, 'title', 'Untitled Section')
            section_content = getattr(section, 'content', '')
            
            # Format the content for better Excel display
            formatted_content = CustomReportReportService.format_content_for_excel(section_content)
            content[title] = formatted_content
            
        return content

    @staticmethod
    def extract_section_order(sections) -> list:
        """Extract section headings in order"""
        order = []
        if not sections:
            return order
            
        for section in sections:
            title = getattr(section, 'title', None)
            if title:
                order.append(title)
        return order

    @staticmethod
    def generate_report_xlsx_file(reports: List[Dict[str, Any]], temp_path: str, study_type: str = None) -> str:
        """
        Generate an Excel file with all case studies of a specific template.
        Includes sections and content for each case study
        """
        try:
            rows = []
            section_order = []
            
            # Extract section order from the first case study
            if reports:
                first_report = reports[0]
                sections = getattr(first_report, 'sections', [])
                section_order = CustomReportReportService.extract_section_order(sections)
            
            for report in reports:
                # Extract basic metadata using getattr for Pydantic models
                # Helper function to handle timezone-aware datetimes and convert to Singapore timezone
                def safe_datetime_sg(dt):
                    if not dt or dt == 'N/A':
                        return dt
                    
                    if hasattr(dt, 'replace'):
                        # Import Singapore timezone
                        from zoneinfo import ZoneInfo
                        SINGAPORE_TZ = ZoneInfo('Asia/Singapore')
                        
                        # Convert to Singapore timezone first
                        if dt.tzinfo is not None:
                            # Convert timezone-aware datetime to Singapore time
                            sg_time = dt.astimezone(SINGAPORE_TZ)
                            # Remove timezone for Excel compatibility
                            return sg_time.replace(tzinfo=None)
                        else:
                            # Assume naive datetime is already in local timezone
                            return dt
                    return dt
                
                display_template_name = study_type

                row = {
                    'report_id': str(getattr(report, '_id', getattr(report, 'id', 'N/A'))),
                    'title': str(getattr(report, 'title', 'N/A')),
                    'template_name': display_template_name,
                    'created_at': safe_datetime_sg(getattr(report, 'created_at', 'N/A')),
                    'updated_at': safe_datetime_sg(getattr(report, 'updated_at', 'N/A')),
                    'user_id': str(getattr(report, 'user_id', 'N/A'))
                }
                
                # Extract section content
                sections = getattr(report, 'sections', [])
                section_content = CustomReportReportService.extract_section_content(sections)
                
                # Add section content to row
                row.update(section_content)
                
                rows.append(row)

            # Get all unique sections
            all_sections = set()
            for row in rows:
                all_sections.update(row.keys())
            
            # Define metadata columns
            meta_cols = ['title', 'template_name', 'report_id', 'user_id', 'created_at', 'updated_at']
            
            # Use the section order from the first case study, then any extra sections
            section_cols = [s for s in section_order if s in all_sections and s not in meta_cols]
            # Add any extra sections not in the first case study at the end
            section_cols += [s for s in all_sections if s not in meta_cols and s not in section_cols]
            columns = meta_cols + section_cols

            # Create DataFrame
            df = pd.DataFrame(rows, columns=columns)

            # Create Excel workbook
            wb = Workbook()
            ws = wb.active
            
            # Set worksheet title
            ws.title = "Reports"

            # Add data to worksheet
            for r_idx, r in enumerate(dataframe_to_rows(df, index=False, header=True)):
                ws.append(r)

            # Style the header row
            header_font = Font(bold=True, color='FFFFFF')
            header_fill = PatternFill(start_color='366092', end_color='366092', fill_type='solid')
            header_alignment = Alignment(horizontal='center', vertical='center')
            
            for cell in ws[1]:
                cell.font = header_font
                cell.fill = header_fill
                cell.alignment = header_alignment

            # Add alternating row colors
            alt_fill = PatternFill(start_color='F2F2F2', end_color='F2F2F2', fill_type='solid')
            for r_idx, row in enumerate(ws.iter_rows(min_row=2), start=2):
                if r_idx % 2 == 1:
                    for cell in row:
                        cell.fill = alt_fill

            # Set text wrapping and alignment for all cells
            for row in ws.iter_rows():
                for cell in row:
                    cell.alignment = Alignment(wrap_text=True, vertical='top')

            # Auto-adjust column widths
            for col in ws.columns:
                col_letter = col[0].column_letter
                header = col[0].value
                
                if header and header not in meta_cols:
                    ws.column_dimensions[col_letter].width = 80
                elif header in ['title']:
                    ws.column_dimensions[col_letter].width = 30
                elif header in ['template_name']:
                    ws.column_dimensions[col_letter].width = 30
                elif header in ['report_id', 'user_id']:
                    ws.column_dimensions[col_letter].width = 25
                elif header in ['created_at', 'updated_at']:
                    ws.column_dimensions[col_letter].width = 20
                else:
                    # Auto-size other metadata columns
                    max_length = max(len(str(cell.value)) if cell.value else 0 for cell in col)
                    ws.column_dimensions[col_letter].width = min(max_length + 2, 50)

            # Save the workbook
            wb.save(temp_path)
            return temp_path

        except Exception as e:
            raise Exception(f"Failed to generate case study Excel file: {str(e)}")