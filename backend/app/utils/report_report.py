from pathlib import Path
import tempfile
import os
import re
import json
import subprocess
import shutil
from typing import Dict, Any, Optional
from jinja2 import Environment, FileSystemLoader
from bson import ObjectId
from docx import Document
from docx.shared import Pt

class ReportReportService:
    """Service for generating and exporting case study PDFs and DOCX"""

    def chr_offset(value):
        """Convert number to letter (0=a, 1=b, etc.)"""
        return chr(ord('a') + value)

    def nl2br(value):
        """Convert newlines to HTML breaks with proper spacing"""
        if not value:
            return ""
        # Replace double newlines (paragraph breaks) with </p><p>
        value = re.sub(r'\n\s*\n', '</p><p>', str(value))
        # Replace single newlines with <br>
        value = re.sub(r'([^\n])\n([^\n])', r'\1<br>\2', value)
        # Wrap in paragraph tags if not already wrapped
        if not value.startswith('<p>'):
            value = f'<p>{value}</p>'
        return value
    
    def clean_br_tags(value):
        """Clean up excessive br tags and spacing"""
        if not value:
            return ""
        # Remove multiple consecutive br tags
        value = re.sub(r'(<br\s*/?>\s*){3,}', '<br><br>', str(value))
        # Remove br tags at the start and end of paragraphs
        value = re.sub(r'<p>\s*<br\s*/?>', '<p>', value)
        value = re.sub(r'<br\s*/?>\s*</p>', '</p>', value)
        # Clean up whitespace around br tags
        value = re.sub(r'\s*<br\s*/?>\s*', '<br>', value)
        return value
    

    # Setup Jinja2 environment for templates
    template_dir = Path(__file__).parent.parent / "templates"
    env = Environment(loader=FileSystemLoader(str(template_dir)))
    env.filters['chr_offset'] = chr_offset
    env.filters['nl2br'] = nl2br
    env.filters['clean_br_tags'] = clean_br_tags

    @staticmethod
    def generate_pdf_report(report_data: Dict[str, Any], output_path: Optional[str] = None) -> str:
        """Generate PDF using LibreOffice conversion from DOCX"""
        
        temp_docx_path = None
        try:
            with tempfile.NamedTemporaryFile(suffix='.docx', delete=False) as temp_docx:
                temp_docx_path = temp_docx.name
            
            ReportReportService.generate_docx_report(
                report_data, temp_docx_path
            )
            
            if not output_path:
                temp_dir = tempfile.gettempdir()
                output_path = os.path.join(temp_dir, f"report_{ObjectId()}.pdf")
            
            output_dir = os.path.dirname(output_path)
            
            print("Converting DOCX to PDF using LibreOffice...")
            
            result = subprocess.run([
                'soffice',
                '--headless',
                '--convert-to', 'pdf',
                '--outdir', output_dir,
                temp_docx_path
            ], 
            capture_output=True, 
            text=True,
            timeout=60 
            )
            
            if result.returncode != 0:
                raise RuntimeError(f"LibreOffice conversion failed: {result.stderr}")
            
            temp_docx_name = os.path.basename(temp_docx_path)
            temp_pdf_name = temp_docx_name.replace('.docx', '.pdf')
            temp_pdf_path = os.path.join(output_dir, temp_pdf_name)
            
            if temp_pdf_path != output_path:
                shutil.move(temp_pdf_path, output_path)
            
            if not os.path.exists(output_path):
                raise RuntimeError("PDF file was not created successfully")
            
            print(f"✅ PDF generated successfully using LibreOffice: {output_path}")
            return output_path
            
        except subprocess.TimeoutExpired:
            raise RuntimeError("LibreOffice conversion timed out after 60 seconds")
        except FileNotFoundError:
            raise RuntimeError("LibreOffice not found. Please install LibreOffice.")
        except Exception as e:
            raise RuntimeError(f"LibreOffice conversion failed: {str(e)}")
        finally:
            if temp_docx_path and os.path.exists(temp_docx_path):
                os.unlink(temp_docx_path)

    @staticmethod
    def generate_docx_report(report_data, output_path):
        # If report_data is a Pydantic model, convert to dict
        if hasattr(report_data, "dict"):
            report_data = report_data.dict()
        doc = Document()

        # Setting the default style
        style = doc.styles['Normal']
        font = style.font
        font.name = 'Times New Roman'
        font.size = Pt(12)

        doc.add_heading(report_data.get("title", "Report"), 0)
        for section in report_data.get("sections", []):
            doc.add_heading(section.get("title", ""), level=1)
            doc.add_paragraph(section.get("content", ""))
        doc.save(output_path)
        return output_path
