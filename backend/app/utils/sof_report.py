from pathlib import Path
import tempfile
import os
import re
import json
import subprocess
import shutil
from datetime import datetime
from typing import Dict, Any, Optional
from jinja2 import Environment, FileSystemLoader
from bson import ObjectId
from docx import Document
from docx.shared import Inches, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH

class SOFReportService:
    """Service for generating and exporting Statement of Facts (Legal) PDFs and DOCX"""

    @staticmethod
    def format_legal_content(value):
        """Format content specifically for legal documents"""
        if not value:
            return ""
        
        # Clean the content first
        content = str(value).strip()
        
        # Remove backslashes pattern like \\\\\\\\\\\\\\\\\\\\\\\
        content = re.sub(r'\\{3,}', '', content)
        
        # Remove ONLY the duplicate headers that appear at the beginning (not the content)
        # Be more specific to avoid removing valid content
        header_patterns = [
            r'^IN THE STATE COURTS OF THE REPUBLIC OF SINGAPORE\s*PUBLIC PROSECUTOR\s*v\s*[A-Z\s]+\s*STATEMENT OF FACTS.*?\n',
        ]
        
        for pattern in header_patterns:
            content = re.sub(pattern, '', content, flags=re.IGNORECASE | re.DOTALL)
        
        # Remove Investigation Officer section from content ONLY at the end (will be handled separately)
        # Be more specific to only remove the signature block, not references to IO in the content
        io_signature_patterns = [
            r'\n\s*(SSGT|SGT|CPL|PC|INSP)\s+[A-Z\s]+BIN\s+[A-Z]+\s*\n\s*INVESTIGATION OFFICER\s*\n\s*SINGAPORE\s*\n\s*\d{1,2}\s+\w+\s+\d{4}\s*$',
            r'\n\s*(SSGT|SGT|CPL|PC|INSP)\s+[A-Z\s]+\s*\n\s*INVESTIGATION OFFICER\s*\n\s*SINGAPORE\s*\n\s*\d{1,2}\s+\w+\s+\d{4}\s*$',
            r'\n\s*DATED this.*?For the Public Prosecutor.*?Singapore \d+\s*$',
            # Additional patterns to catch mid-content IO signatures
            r'\n\s*(SSGT|SGT|CPL|PC|INSP)\s+[A-Z\s]+(?:BIN\s+[A-Z]+)?\s*\n\s*INVESTIGATION OFFICER\s*(?:\n|$)',
            r'\n\s*INVESTIGATION OFFICER\s*\n\s*SINGAPORE\s*(?:\n|$)',
            r'(SSGT|SGT|CPL|PC|INSP)\s+[A-Z\s]+(?:BIN\s+[A-Z]+)?\s+INVESTIGATION OFFICER(?:\s+SINGAPORE)?(?:\s+\d{1,2}\s+\w+\s+\d{4})?',
        ]
        
        for pattern in io_signature_patterns:
            content = re.sub(pattern, '', content, flags=re.IGNORECASE | re.DOTALL)
        
        # Remove any existing HTML tags except tables
        content = re.sub(r'<(?!/?(?:table|tr|td|th|thead|tbody|br)\b)[^>]*>', '', content)
        
        # Handle "subheader:" prefix and mark for later processing
        content = re.sub(
            r'^subheader:\s*(.+)$',
            r'SUBHEADER_MARKER:\1',
            content,
            flags=re.MULTILINE | re.IGNORECASE
        )
        
        # Format main section headers
        content = re.sub(
            r'(FACTS RELATING TO THE PROCEEDED\s+([A-Z]+)\s+CHARGE\s*[-–]\s*THE\s+([A-Z0-9]+)\s+CHARGE\s*\([^)]+\))',
            r'<strong>\1</strong>',
            content,
            flags=re.IGNORECASE
        )
        
        # Handle numbered paragraphs with proper line breaks
        # Split content into paragraphs first
        paragraphs = content.split('\n')
        formatted_paragraphs = []
        
        for para in paragraphs:
            para = para.strip()
            if not para:
                continue
            
            # ONLY skip IO signature blocks that appear at the end, not content references
            # Be very specific about what to skip
            if (para.upper().startswith('SSGT ') or para.upper().startswith('SGT ') or 
                para.upper().startswith('CPL ') or para.upper().startswith('PC ') or 
                para.upper().startswith('INSP ')) and \
            ('INVESTIGATION OFFICER' in content[content.find(para):content.find(para)+200] or
             len(para.split()) <= 4):  # Also skip if it's just the name/rank line
                continue
            if para.upper() == 'INVESTIGATION OFFICER':
                continue
            if para.upper() == 'SINGAPORE' and len(para) < 20:  # Only skip standalone "SINGAPORE", not in sentences
                continue
            if re.match(r'^\d{1,2}\s+\w+\s+\d{4}$', para):  # Only skip standalone dates
                continue
            if 'DATED THIS' in para.upper() and 'PUBLIC PROSECUTOR' in para.upper():
                continue
            # Skip lines that are just IO names without context (likely duplicates)
            if re.match(r'^(SSGT|SGT|CPL|PC|INSP)\s+[A-Z\s]+$', para.upper()) and len(para.split()) <= 4:
                continue
                
            # Check if it's a numbered paragraph (starts with number followed by period)
            if re.match(r'^\d+\.', para):
                # Add line break before numbered paragraphs (except first one)
                if formatted_paragraphs:
                    formatted_paragraphs.append('<br>')
                formatted_paragraphs.append(f'<p class="numbered-paragraph">{para}</p>')
            
            # Check if it's a roman numeral sub-point (comprehensive pattern: i-xx) - check this BEFORE lettered
            elif re.match(r'^(i{1,3}|iv|v|vi{1,3}|ix|x|xi{1,3}|xiv|xv|xvi{1,3}|xix|xx)\.', para):
                formatted_paragraphs.append(f'<p class="roman-paragraph">{para}</p>')

            # Check if it's a lettered sub-point (can now include all letters since roman is checked first)
            elif re.match(r'^[a-z]\.', para):
                formatted_paragraphs.append(f'<p class="lettered-paragraph">{para}</p>')
            
            # Check if it's a parenthesized roman numeral sub-point (4th level: (i), (ii), etc.)
            elif re.match(r'^\((i{1,3}|iv|v|vi{1,3}|ix|x|xi{1,3}|xiv|xv|xvi{1,3}|xix|xx)\)', para):
                formatted_paragraphs.append(f'<p class="parenthesis-roman-paragraph">{para}</p>')
            
            # Check if it's a section header
            elif para.startswith('FACTS RELATING') or para.startswith('<strong>FACTS'):
                if formatted_paragraphs:
                    formatted_paragraphs.append('<br>')
                formatted_paragraphs.append(f'<p class="section-header">{para}</p>')
                formatted_paragraphs.append('<br>')
            
            # Check if it's a subheading (italic)
            elif para.startswith('<em>') and para.endswith('</em>'):
                formatted_paragraphs.append('<br>')
                formatted_paragraphs.append(f'<p class="subheading">{para}</p>')
            
            # Check if it's a subheader (from "subheader:" prefix)
            elif para.startswith('SUBHEADER_MARKER:'):
                clean_text = para.replace('SUBHEADER_MARKER:', '').strip()
                formatted_paragraphs.append('<br>')
                formatted_paragraphs.append(f'<p class="subheading">{clean_text}</p>')
            
            # Check if it's a table
            elif '<table' in para:
                formatted_paragraphs.append(f'<div class="table-container">{para}</div>')
            
            # Regular paragraph
            else:
                formatted_paragraphs.append(f'<p class="regular-paragraph">{para}</p>')
        
        content = '\n'.join(formatted_paragraphs)
        
        # Clean up excessive line breaks
        content = re.sub(r'(<br>\s*){3,}', '<br><br>', content)
        
        return content

    # Setup Jinja2 environment for templates
    template_dir = Path(__file__).parent.parent / "templates"
    env = Environment(loader=FileSystemLoader(str(template_dir)))
    env.filters['format_legal_content'] = format_legal_content

    @staticmethod
    def generate_legal_statement_pdf(report_data: Dict[str, Any], output_path: Optional[str] = None) -> str:
        """Generate legal Statement of Facts PDF using LibreOffice conversion from DOCX"""
        
        temp_docx_path = None
        try:
            with tempfile.NamedTemporaryFile(suffix='.docx', delete=False) as temp_docx:
                temp_docx_path = temp_docx.name
            
            SOFReportService.generate_legal_statement_docx(
                report_data, temp_docx_path
            )
            
            if not output_path:
                temp_dir = tempfile.gettempdir()
                output_path = os.path.join(temp_dir, f"legal_statement_{ObjectId()}.pdf")
            
            output_dir = os.path.dirname(output_path)
            
            print("Converting legal statement DOCX to PDF using LibreOffice...")
            
            result = subprocess.run([
                'soffice',
                '--headless',
                '--convert-to', 'pdf',
                '--outdir', output_dir,
                temp_docx_path
            ], 
            capture_output=True, 
            text=True,
            timeout=60 
            )
            
            if result.returncode != 0:
                raise RuntimeError(f"LibreOffice conversion failed: {result.stderr}")
            
            temp_docx_name = os.path.basename(temp_docx_path)
            temp_pdf_name = temp_docx_name.replace('.docx', '.pdf')
            temp_pdf_path = os.path.join(output_dir, temp_pdf_name)
            
            if temp_pdf_path != output_path:
                shutil.move(temp_pdf_path, output_path)
            
            if not os.path.exists(output_path):
                raise RuntimeError("PDF file was not created successfully")
            
            print(f"✅ Legal statement PDF generated successfully using LibreOffice: {output_path}")
            return output_path
            
        except subprocess.TimeoutExpired:
            raise RuntimeError("LibreOffice conversion timed out after 60 seconds")
        except FileNotFoundError:
            raise RuntimeError("LibreOffice not found. Please install LibreOffice.")
        except Exception as e:
            raise RuntimeError(f"LibreOffice conversion failed: {str(e)}")
        finally:
            if temp_docx_path and os.path.exists(temp_docx_path):
                os.unlink(temp_docx_path)

    @staticmethod
    def generate_legal_statement_docx(report_data: Dict[str, Any], output_path: Optional[str] = None) -> str:
        """Generate a legal Statement of Facts DOCX from case study data"""
        
        # Process the case study data into legal format
        legal_data = SOFReportService._process_legal_statement_data(report_data)
        generation_date = datetime.now().strftime("%d %B %Y")
        
        # Determine output path
        if not output_path:
            temp_dir = tempfile.gettempdir()
            output_path = os.path.join(temp_dir, f"{ObjectId()}.docx")

        # Create a new Document
        document = Document()
        
        # Set default font and size via "Normal" style
        style = document.styles['Normal']
        font = style.font
        font.name = 'Times New Roman'
        font.size = Pt(12)
        
        # Set document margins (25mm top/bottom, 20mm left/right)
        sections = document.sections
        for section in sections:
            section.top_margin = 914400      # 2.54 cm
            section.bottom_margin = 914400   # 2.54 cm
            section.left_margin = 1144800    # 3.18 cm
            section.right_margin = 1144800   # 3.18 cm
        number_of_underlines = 71 # Adjust accordingly with the margins

        # Set default font to Times New Roman 12pt
        from docx.oxml.shared import OxmlElement, qn

        # Add document header
        header_p = document.add_paragraph()
        SOFReportService._set_line_spacing(header_p)
        header_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        header_run = header_p.add_run("IN THE STATE COURTS OF THE REPUBLIC OF SINGAPORE")
        header_run.font.bold = True
        header_run.underline = True
        
        # PUBLIC PROSECUTOR
        pp_p = document.add_paragraph()
        SOFReportService._set_line_spacing(pp_p)
        pp_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        pp_run = pp_p.add_run("PUBLIC PROSECUTOR")
        pp_run.font.bold = True

        # v
        v_p = document.add_paragraph()
        SOFReportService._set_line_spacing(v_p)
        v_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        v_run = v_p.add_run("v")
        v_run.font.bold = True

        # ACCUSED NAME
        accused_p = document.add_paragraph()
        SOFReportService._set_line_spacing(accused_p)
        accused_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        accused_run = accused_p.add_run(legal_data.get('accused_name', 'ACCUSED'))
        accused_run.font.bold = True

        # Add horizontal line
        line_one = document.add_paragraph()
        SOFReportService._set_line_spacing(line_one)
        run = line_one.add_run("_" * number_of_underlines)
        run.bold = True

        # Document title
        title_p = document.add_paragraph()
        SOFReportService._set_line_spacing(title_p)
        title_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        title_p.paragraph_format.space_after = Pt(0)    # Removes the space after the paragraph
        title_run = title_p.add_run(f"STATEMENT OF FACTS (FOR {legal_data.get('charge_count')} PROCEEDED CHARGES)")
        title_run.font.bold = True

        # Add horizontal line
        line_two = document.add_paragraph()
        SOFReportService._set_line_spacing(line_two)
        run = line_two.add_run("_" * number_of_underlines)
        run.bold = True
        
        # Add main content
        statement_content = legal_data.get('statement_content', '')
        
        content_paragraphs = SOFReportService._format_legal_content_for_docx(statement_content)
        
        for para_data in content_paragraphs:
            if para_data['type'] == 'numbered':
                p = document.add_paragraph(style='List Number')
                SOFReportService._set_line_spacing(p)
                p.paragraph_format.space_after = Pt(0)    # Removes the space after the paragraph
                p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
                run = p.add_run(para_data['text'])
                                
            elif para_data['type'] == 'lettered':
                p = document.add_paragraph()
                SOFReportService._set_line_spacing(p)
                p.paragraph_format.space_after = Pt(0)    # Removes the space after the paragraph
                p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
                p.paragraph_format.left_indent = Inches(0.5)  # 0.5 inch indent
                p.paragraph_format.first_line_indent = Inches(-0.25)  # Hanging indent
                run = p.add_run(para_data['text'])
                
            elif para_data['type'] == 'roman':
                p = document.add_paragraph()
                SOFReportService._set_line_spacing(p)
                p.paragraph_format.space_after = Pt(0)    # Removes the space after the paragraph
                p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
                p.paragraph_format.left_indent = Inches(1.0)  # 1 inch indent
                p.paragraph_format.first_line_indent = Inches(-0.25)  # Hanging indent (same as lettered)
                run = p.add_run(para_data['text'])
                
            elif para_data['type'] == 'parenthesis_roman':
                p = document.add_paragraph()
                SOFReportService._set_line_spacing(p)
                p.paragraph_format.space_after = Pt(0)    # Removes the space after the paragraph
                p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
                p.paragraph_format.left_indent = Inches(1.5)  # 1.5 inch indent (deepest level)
                p.paragraph_format.first_line_indent = Inches(-0.25)  # Hanging indent
                run = p.add_run(para_data['text'])
                
            elif para_data['type'] == 'section_header':
                document.add_paragraph()
                p = document.add_paragraph()
                SOFReportService._set_line_spacing(p)
                p.alignment = WD_ALIGN_PARAGRAPH.LEFT
                run = p.add_run(para_data['text'])
                run.font.bold = True
                run.underline = True
                
            elif para_data['type'] == 'subheading':
                document.add_paragraph()
                p = document.add_paragraph()
                SOFReportService._set_line_spacing(p)
                p.alignment = WD_ALIGN_PARAGRAPH.LEFT
                run = p.add_run(para_data['text'])
                run.font.italic = True
                
            elif para_data['type'] == 'table':
                # Handle table content properly
                SOFReportService._add_table_to_docx(document, para_data['table_data'])
                
            else:  # regular paragraph
                p = document.add_paragraph()
                SOFReportService._set_line_spacing(p)
                p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
                run = p.add_run(para_data['text'])
            
        # Add signature section if IO details are available
        if legal_data.get('io_name') and legal_data.get('io_name') != 'INVESTIGATION OFFICER':
            # Add some space before signature
            document.add_paragraph()
            document.add_paragraph()
            document.add_paragraph()
            
            # IO Name
            io_name_p = document.add_paragraph()
            io_name_p.paragraph_format.space_after = Pt(0)    # Removes the space after the paragraph
            io_name_run = io_name_p.add_run(f"{legal_data.get('io_rank', '')} {legal_data.get('io_name', '')}".strip())
            io_name_run.font.bold = True

            # IO Title
            io_title_p = document.add_paragraph()
            SOFReportService._set_line_spacing(io_title_p)
            io_title_p.paragraph_format.space_after = Pt(0)    # Removes the space after the paragraph
            io_title_run = io_title_p.add_run("INVESTIGATION OFFICER")
            io_title_run.font.bold = True

            # Location
            location_p = document.add_paragraph()
            SOFReportService._set_line_spacing(location_p)
            location_p.paragraph_format.space_after = Pt(0)    # Removes the space after the paragraph
            location_run = location_p.add_run("SINGAPORE")
            location_run.font.bold = True

            # Date
            date_p = document.add_paragraph()
            date_run = date_p.add_run(generation_date)

        # Save the document
        document.save(output_path)
        return output_path

    @staticmethod
    def _format_legal_content_for_docx(content):
        """Format legal content for DOCX with proper paragraph types and table handling"""
        
        # First, clean the content\
        content = str(content).strip()
        
        # Remove Investigation Officer signature patterns from content
        io_signature_patterns = [
            r'\n\s*(SSGT|SGT|CPL|PC|INSP)\s+[A-Z\s]+BIN\s+[A-Z]+\s*\n\s*INVESTIGATION OFFICER\s*\n\s*SINGAPORE\s*\n\s*\d{1,2}\s+\w+\s+\d{4}\s*$',
            r'\n\s*(SSGT|SGT|CPL|PC|INSP)\s+[A-Z\s]+\s*\n\s*INVESTIGATION OFFICER\s*\n\s*SINGAPORE\s*\n\s*\d{1,2}\s+\w+\s+\d{4}\s*$',
            r'\n\s*DATED this.*?For the Public Prosecutor.*?Singapore \d+\s*$',
        ]
        
        for pattern in io_signature_patterns:
            content = re.sub(pattern, '', content, flags=re.IGNORECASE | re.DOTALL)
        
        paragraphs = []
        lines = content.split('\n')
        
        i = 0
        while i < len(lines):
            line = lines[i].strip()
            if not line:
                i += 1
                continue
            
            # Skip IO signature blocks that appear in the content
            if (line.upper().startswith('SSGT ') or line.upper().startswith('SGT ') or 
                line.upper().startswith('CPL ') or line.upper().startswith('PC ') or
                line.upper().startswith('INSP ')) and \
                i + 1 < len(lines) and 'INVESTIGATION OFFICER' in lines[i + 1].upper():
                # Skip this line and the next few lines that are part of signature
                while i < len(lines) and (
                    lines[i].strip().upper().startswith(('SSGT', 'SGT', 'CPL', 'PC', 'INSP')) or
                    'INVESTIGATION OFFICER' in lines[i].upper() or
                    lines[i].strip().upper() == 'SINGAPORE' or
                    re.match(r'^\d{1,2}\s+\w+\s+\d{4}$', lines[i].strip())
                ):
                    i += 1
                continue
                
            if line.upper() == 'INVESTIGATION OFFICER':
                i += 1
                continue
            if line.upper() == 'SINGAPORE' and len(line) < 20:  # Only skip standalone "SINGAPORE"
                i += 1
                continue
            if re.match(r'^\d{1,2}\s+\w+\s+\d{4}$', line):  # Only skip standalone dates
                i += 1
                continue
            if 'DATED THIS' in line.upper() and 'PUBLIC PROSECUTOR' in line.upper():
                i += 1
                continue
            
            # Check if this line starts a table
            if '<table' in line:
                # Extract the complete table
                table_html = ""
                while i < len(lines) and '</table>' not in lines[i]:
                    table_html += lines[i] + '\n'
                    i += 1
                if i < len(lines):
                    table_html += lines[i]  # Add the closing </table>
                    i += 1
                
                # Parse the table
                table_data = SOFReportService._parse_html_table(table_html)
                if table_data:
                    paragraphs.append({'type': 'table', 'table_data': table_data})
                continue
                
            # Check paragraph type based on patterns
            if re.match(r'^\d+\.', line):  # Numbered paragraph
                cleaned_line = re.sub(r'^\d+\.\s*', '', line)
                paragraphs.append({'type': 'numbered', 'text': cleaned_line})
            elif re.match(r'^(i{1,3}|iv|v|vi{1,3}|ix|x|xi{1,3}|xiv|xv|xvi{1,3}|xix|xx)\.', line):  # Roman numeral paragraph (comprehensive 1-20) - check BEFORE lettered
                paragraphs.append({'type': 'roman', 'text': line})
            elif re.match(r'^[a-z]\.', line):  # Lettered paragraph (now can include all letters)
                paragraphs.append({'type': 'lettered', 'text': line})
            elif re.match(r'^\((i{1,3}|iv|v|vi{1,3}|ix|x|xi{1,3}|xiv|xv|xvi{1,3}|xix|xx)\)', line):  # Parenthesized roman numeral (4th level)
                paragraphs.append({'type': 'parenthesis_roman', 'text': line})
            elif line.isupper() and len(line) > 5:  # Section headers (all caps)
                paragraphs.append({'type': 'section_header', 'text': line})
            elif re.match(r'^subheader:\s*', line, re.IGNORECASE):  # Subheadings using "subheader:" prefix
                clean_text = re.sub(r'^subheader:\s*', '', line, flags=re.IGNORECASE).strip()
                paragraphs.append({'type': 'subheading', 'text': clean_text})
            else:  # Regular paragraph
                paragraphs.append({'type': 'regular', 'text': line})
            
            i += 1
                
        return paragraphs

    @staticmethod
    def _parse_html_table(table_html):
        """Parse HTML table into structured data for DOCX table creation"""
        import re
        
        # Extract headers
        header_pattern = r'<th[^>]*>(.*?)</th>'
        headers = re.findall(header_pattern, table_html, re.DOTALL | re.IGNORECASE)
        headers = [re.sub(r'<[^>]*>', '', header).strip() for header in headers]
        
        # Extract rows
        row_pattern = r'<tr[^>]*>(.*?)</tr>'
        rows_html = re.findall(row_pattern, table_html, re.DOTALL | re.IGNORECASE)
        
        rows = []
        for row_html in rows_html:
            # Skip header rows
            if '<th' in row_html:
                continue
                
            cell_pattern = r'<td[^>]*>(.*?)</td>'
            cells = re.findall(cell_pattern, row_html, re.DOTALL | re.IGNORECASE)
            # Clean HTML tags and normalize whitespace
            cells = [re.sub(r'<[^>]*>', ' ', cell).strip() for cell in cells]
            cells = [re.sub(r'\s+', ' ', cell) for cell in cells]
            if cells:
                rows.append(cells)
        
        if headers and rows:
            return {'headers': headers, 'rows': rows}
        return None

    @staticmethod
    def _add_table_to_docx(document, table_data):
        """Add a properly formatted table to the DOCX document"""
        if not table_data or 'headers' not in table_data or 'rows' not in table_data:
            return
        
        headers = table_data['headers']
        rows = table_data['rows']
        
        # Create table with proper number of rows and columns
        table = document.add_table(rows=1 + len(rows), cols=len(headers))
        table.style = 'Table Grid'
        
        # Set header row
        header_row = table.rows[0]
        for i, header in enumerate(headers):
            cell = header_row.cells[i]
            cell.text = header
            # Format header cell
            paragraph = cell.paragraphs[0]
            SOFReportService._set_line_spacing(paragraph)
            run = paragraph.runs[0] if paragraph.runs else paragraph.add_run(header)
            paragraph.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
            
            # Add cell borders and shading
            from docx.oxml.shared import OxmlElement, qn
            from docx.oxml.ns import nsdecls
            
            # Add shading to header
            shading_elm = OxmlElement('w:shd')
            shading_elm.set(qn('w:fill'), 'D9D9D9')  # Light gray background
            cell._tc.get_or_add_tcPr().append(shading_elm)
        
        # Add data rows
        for row_idx, row_data in enumerate(rows):
            table_row = table.rows[row_idx + 1]
            for col_idx, cell_data in enumerate(row_data):
                if col_idx < len(table_row.cells):
                    cell = table_row.cells[col_idx]
                    cell.text = str(cell_data)
                    # Format data cell
                    paragraph = cell.paragraphs[0]
                    SOFReportService._set_line_spacing(paragraph)
                    if paragraph.runs:
                        run = paragraph.runs[0]
                    else:
                        run = paragraph.add_run(str(cell_data))
                    paragraph.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        
        # Set column widths for better appearance
        if len(headers) == 4:  # Standard evidence table format
            widths = [Inches(1.5), Inches(0.8), Inches(2.5), Inches(2.2)]  # Adjust as needed
            for i, width in enumerate(widths):
                if i < len(table.columns):
                    for cell in table.columns[i].cells:
                        cell.width = width

    @staticmethod
    def _extract_charge_count(content):
        """Extract the number of proceeded charges based on header occurrences"""
        matches = re.findall(r'FACTS\s+RELATING\s+TO\s+THE\s+PROCEEDED', content, re.IGNORECASE)
        
        if matches:
            return str(len(matches))
        
        # Fallback: count DAC references
        dac_matches = re.findall(r'DAC-\d+-\d+', content)
        if dac_matches:
            return str(len(set(dac_matches)))
        
        # Default
        return "2"

    @staticmethod
    def _process_legal_statement_data(report_data):
        """Process case study data into legal statement format"""
        
        # Extract and process the main content
        statement_content = ""
        for section in report_data.get('sections', []):
            statement_content += section.get('content', '')
        
        # Extract charge count
        charge_count = SOFReportService._extract_charge_count(statement_content)
        
        # Extract accused details with better parsing
        accused_details = SOFReportService._extract_accused_details_enhanced(statement_content)
        arrest_details = SOFReportService._extract_arrest_details_enhanced(statement_content)
        exhibits = SOFReportService._extract_exhibits_enhanced(statement_content)
        analysis_results = SOFReportService._extract_analysis_results_enhanced(statement_content)
        charges = SOFReportService._extract_charges_enhanced(statement_content)
        
        # Extract IO details including signing date
        io_details = SOFReportService._extract_io_details(statement_content)
        
        return {
            'case_title': report_data.get('title', 'STATEMENT OF FACTS'),
            'accused_name': accused_details.get('name', 'ACCUSED NAME'),
            'accused_details': accused_details,
            'arrest_details': arrest_details,
            'exhibits': exhibits,
            'analysis_results': analysis_results,
            'charges': charges,
            'charge_count': charge_count,
            'io_name': io_details.get('name', 'INVESTIGATION OFFICER'),
            'io_rank': io_details.get('rank', 'SSGT'),
            'statement_content': statement_content,
            'sections': report_data.get('sections', [])
        }

    @staticmethod
    def _extract_accused_details_enhanced(content):
        """Enhanced extraction of accused person details"""
        details = {'name': 'ACCUSED NAME', 'age': '', 'nric': '', 'occupation': '', 'nationality': ''}
        
        # Extract name pattern
        name_match = re.search(r'accused is ([A-Z\s]+)(?:\s*\(|,)', content, re.IGNORECASE)
        if name_match:
            details['name'] = name_match.group(1).strip().upper()
        
        # Extract NRIC
        nric_match = re.search(r'NRIC No\.?:?\s*([SsTtFfGg]\d{7}[A-Za-z])', content)
        if nric_match:
            details['nric'] = nric_match.group(1).upper()
        
        # Extract age
        age_match = re.search(r'(\d{1,2})[-\s]*year[-\s]*old', content, re.IGNORECASE)
        if age_match:
            details['age'] = age_match.group(1)
        
        # Extract nationality
        nationality_match = re.search(r'year[-\s]*old\s+(\w+)', content, re.IGNORECASE)
        if nationality_match:
            details['nationality'] = nationality_match.group(1)
        
        # Extract occupation
        occupation_match = re.search(r'was\s+(working as|employed as)?\s*([^.]+)\s+at the material time', content, re.IGNORECASE)
        if occupation_match:
            details['occupation'] = occupation_match.group(2).strip()
        
        return details

    @staticmethod
    def _extract_arrest_details_enhanced(content):
        """Enhanced extraction of arrest details"""
        details = {'date': '', 'time': '', 'location': ''}
        
        # Extract date
        date_match = re.search(r'On\s+(\d{1,2}\s+\w+\s+\d{4})', content)
        if date_match:
            details['date'] = date_match.group(1)
        
        # Extract time
        time_match = re.search(r'at about\s+(\d{1,2}\.\d{2}[ap]m|\d{1,2}:\d{2}[ap]m)', content, re.IGNORECASE)
        if time_match:
            details['time'] = time_match.group(1)
        
        # Extract location
        location_match = re.search(r'at\s+(Blk\s+\d+[^,]+)', content)
        if location_match:
            details['location'] = location_match.group(1)
        
        return details

    @staticmethod
    def _extract_exhibits_enhanced(content):
        """Enhanced extraction of exhibits"""
        exhibits = []
        
        # Find exhibit markings and descriptions
        exhibit_patterns = [
            r'(One\s+[^(]+)\s*\(later marked as\s*["\']([A-Z0-9]+)["\']',
            r'([^(]+)\s*\(marked as\s*["\']([A-Z0-9]+)["\']'
        ]
        
        for pattern in exhibit_patterns:
            matches = re.findall(pattern, content, re.IGNORECASE)
            for match in matches:
                exhibits.append({
                    'description': match[0].strip(),
                    'marking': match[1]
                })
        
        return exhibits

    @staticmethod
    def _extract_analysis_results_enhanced(content):
        """Enhanced extraction of analysis results"""
        results = []
        
        # Look for analysis results in table format or structured text
        lab_no_pattern = r'(\d+[-\w]+)\s*\((\d{1,2}\s+\w+\s+\d{4})\)'
        weight_pattern = r'not less than\s+([\d.]+g)\s+of\s+(\w+)'
        marking_pattern = r'marked?\s+(?:as\s+)?["\']?([A-Z0-9]+)["\']?'
        
        # Extract lab numbers and dates
        lab_matches = re.findall(lab_no_pattern, content)
        # Extract weights and substances
        weight_matches = re.findall(weight_pattern, content, re.IGNORECASE)
        # Extract markings
        marking_matches = re.findall(marking_pattern, content)
        
        if lab_matches and weight_matches:
            for i, (lab_no, date) in enumerate(lab_matches):
                if i < len(weight_matches):
                    weight, substance = weight_matches[i]
                    marking = marking_matches[i] if i < len(marking_matches) else f'EXHIBIT_{i+1}'
                    
                    results.append({
                        'lab_no': f'{lab_no} ({date})',
                        'marking': marking,
                        'description': f'Substance containing {substance}',
                        'weight': f'Not less than {weight} of {substance}',
                        'substance': substance
                    })
        
        return results

    @staticmethod
    def _extract_charges_enhanced(content):
        """Enhanced extraction of charges"""
        charges = []
        
        # Look for section references
        charge_patterns = [
            r'section\s+(\d+\([a-z]\))\s+of\s+the\s+MDA',
            r'section\s+(\d+\(\d+\)\([a-z]\))\s+read\s+with\s+section\s+(\d+\(\d+\))\s+of\s+the\s+MDA'
        ]
        
        for pattern in charge_patterns:
            matches = re.findall(pattern, content, re.IGNORECASE)
            for match in matches:
                if isinstance(match, tuple):
                    section = f"Section {match[0]} read with Section {match[1]} of the MDA"
                    charge_type = 'trafficking' if '5(' in match[0] else 'possession'
                else:
                    section = f"Section {match} of the MDA"
                    charge_type = 'trafficking' if '5(' in match else 'possession'
                
                charges.append({
                    'type': charge_type,
                    'section': section,
                    'description': f'{charge_type.title()} of controlled drug'
                })
        
        return charges

    @staticmethod
    def _extract_io_details(content):
        """Extract Investigation Officer details including signing date"""
        details = {'name': '', 'rank': '', 'sign_date': ''}
        
        # Look for IO signature section patterns
        io_patterns = [
            r'((?:SSGT|SGT|CPL|PC|INSP)\s+[A-Z\s]+BIN\s+[A-Z]+)\s*INVESTIGATION OFFICER\s*SINGAPORE\s*(\d{1,2}\s+\w+\s+\d{4})',
            r'((?:SSGT|SGT|CPL|PC|INSP)\s+[A-Z\s]+)\s*INVESTIGATION OFFICER\s*SINGAPORE\s*(\d{1,2}\s+\w+\s+\d{4})'
        ]
        
        for pattern in io_patterns:
            match = re.search(pattern, content, re.IGNORECASE | re.DOTALL)
            if match:
                full_name = match.group(1).strip()
                sign_date = match.group(2).strip()
                
                parts = full_name.split(None, 1)
                if len(parts) >= 2:
                    details['rank'] = parts[0].upper()
                    details['name'] = parts[1].upper()
                details['sign_date'] = sign_date
                break
        
        # If not found with date, try without date first
        if not details['name']:
            rank_name_patterns = [
                r'((?:SSGT|SGT|CPL|PC|INSP)\s+[A-Z\s]+BIN\s+[A-Z]+)\s*INVESTIGATION OFFICER',
                r'((?:SSGT|SGT|CPL|PC|INSP)\s+[A-Z\s]+)\s*INVESTIGATION OFFICER'
            ]
            
            for pattern in rank_name_patterns:
                match = re.search(pattern, content, re.IGNORECASE)
                if match:
                    full_name = match.group(1).strip()
                    parts = full_name.split(None, 1)
                    if len(parts) >= 2:
                        details['rank'] = parts[0].upper()
                        details['name'] = parts[1].upper()
                    break
        
        return details

    @staticmethod
    def _extract_accused_details(report_data):
        """Extract accused person details from case study data"""
        details = {'name': 'ACCUSED NAME', 'age': '', 'nric': '', 'occupation': ''}
        
        for section in report_data.get('sections', []):
            content = section.get('content', '').lower()
            if 'accused' in section.get('title', '').lower() or 'accused' in content:
                # Extract NRIC
                nric_match = re.search(r'[sStTfFgG]\d{7}[a-zA-Z]', section.get('content', ''))
                if nric_match:
                    details['nric'] = nric_match.group(0).upper()
                
                # Extract age
                age_match = re.search(r'(\d{1,2})[-\s]*year[-\s]*old', content)
                if age_match:
                    details['age'] = age_match.group(1)
                
                # Extract occupation
                if 'occupation' in content or 'working as' in content:
                    details['occupation'] = 'As stated in case documents'
                    
        return details

    @staticmethod
    def _extract_arrest_details(report_data):
        """Extract arrest details from case study data"""
        details = {'date': '', 'time': '', 'location': '', 'officers': ''}
        
        for section in report_data.get('sections', []):
            content = section.get('content', '')
            if 'arrest' in section.get('title', '').lower() or 'arrest' in content.lower():
                # Extract date patterns
                date_match = re.search(r'(\d{1,2}\s+\w+\s+\d{4})', content)
                if date_match:
                    details['date'] = date_match.group(1)
                
                # Extract time patterns
                time_match = re.search(r'(\d{1,2}[:.]\d{2}\s*[ap]m)', content, re.IGNORECASE)
                if time_match:
                    details['time'] = time_match.group(1)
                    
                # Extract location
                if 'blk' in content.lower() or 'block' in content.lower():
                    details['location'] = 'As stated in case documents'
                    
        return details

    @staticmethod
    def _extract_exhibits(report_data):
        """Extract exhibit information from case study data"""
        exhibits = []
        
        for section in report_data.get('sections', []):
            if 'exhibit' in section.get('title', '').lower() or 'seized' in section.get('title', '').lower():
                content = section.get('content', '')
                
                # Look for exhibit markings
                exhibit_matches = re.findall(r'marked as ["\']([A-Z0-9]+)["\']', content)
                for match in exhibit_matches:
                    exhibits.append({
                        'marking': match,
                        'description': 'As described in case documents'
                    })
                    
        return exhibits

    @staticmethod
    def _extract_analysis_results(report_data):
        """Extract laboratory analysis results"""
        results = []
        
        for section in report_data.get('sections', []):
            if 'analysis' in section.get('title', '').lower():
                content = section.get('content', '')
                
                # Look for controlled substances
                if 'methamphetamine' in content.lower():
                    results.append({
                        'lab_no': 'As per HSA Certificate',
                        'marking': 'As marked',
                        'description': 'Crystalline/powdery substance',
                        'weight': 'As analysed',
                        'substance': 'Methamphetamine'
                    })
                
                if 'cannabis' in content.lower():
                    results.append({
                        'lab_no': 'As per HSA Certificate',
                        'marking': 'As marked',
                        'description': 'Vegetable matter',
                        'weight': 'As analysed',
                        'substance': 'Cannabis'
                    })
                
        return results

    @staticmethod
    def _extract_charges(report_data):
        """Extract charge information"""
        charges = []
        
        for section in report_data.get('sections', []):
            if 'charge' in section.get('title', '').lower():
                content = section.get('content', '')
                
                if 'section 8' in content.lower():
                    charges.append({
                        'type': 'possession',
                        'section': 'Section 8(a) of the MDA',
                        'description': 'Possession of controlled drug'
                    })
                
                if 'section 5' in content.lower():
                    charges.append({
                        'type': 'trafficking',
                        'section': 'Section 5(1)(a) read with Section 5(2) of the MDA',
                        'description': 'Possession for purpose of trafficking'
                    })
                    
        return charges
    
    @staticmethod
    def _set_line_spacing(paragraph, spacing=1.5):
        paragraph.paragraph_format.line_spacing = spacing